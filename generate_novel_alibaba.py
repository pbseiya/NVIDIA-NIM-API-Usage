import os
import time
import re
from openai import OpenAI
from dotenv import load_dotenv
from rich.console import Console
from docx import Document
from docx.shared import RGBColor, Pt

load_dotenv()
# Alibaba DashScope compatible mode API key
api_key = os.getenv("ALIBABA_API_KEY")
# Alibaba DashScope compatible mode base URL
client = OpenAI(base_url="https://coding-intl.dashscope.aliyuncs.com/v1", api_key=api_key)

console = Console()

def style_run(run, bold=False, italic=False, r=0, g=0, b=0, size=14):
    """Helper function to apply styles to a run"""
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.size = Pt(size)
    run.font.name = "Cordia New"

def create_docx(text, filename="Sea_Cat_and_Old_Turtle_Alibaba.docx"):
    doc = Document()

    # Custom Title & Formatting
    title = doc.add_heading("แมวทะเล กับ ผู้เฒ่าเต่า (Alibaba Qwen Version)", 0)
    title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    title.runs[0].font.name = "Cordia New"

    paragraphs = text.split("\n")
    for p in paragraphs:
        if not p.strip():
            continue

        if p.startswith("#"):
            h_level = min(len(p) - len(p.lstrip("#")), 3)
            clean_text = p.strip("# ")
            head = doc.add_heading(clean_text, level=h_level)
            head.runs[0].font.color.rgb = RGBColor(153, 51, 255)
            continue

        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_after = Pt(12)

        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", p)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p_obj.add_run(part[2:-2])
                style_run(r, bold=True, r=204, g=0, b=0)
            elif part.startswith("*") and part.endswith("*"):
                r = p_obj.add_run(part[1:-1])
                style_run(r, italic=True, r=0, g=153, b=0)
            else:
                words = re.split(r"(แมวทะเล|ผู้เฒ่าเต่า)", part)
                for w in words:
                    r = p_obj.add_run(w)
                    if w == "แมวทะเล":
                        style_run(r, bold=True, r=0, g=153, b=255)
                    elif w == "ผู้เฒ่าเต่า":
                        style_run(r, bold=True, r=102, g=51, b=0)
                    else:
                        style_run(r, r=50, g=50, b=50)

    doc.save(filename)
    console.print(f"\n[bold green]✅ บันทึกนิยายลงไฟล์ {filename} เรียบร้อยแล้ว![/bold green]")

def run_novel_generation():
    # Alibaba model name (qwen-plus or qwen-max)
    # The user mentioned qwen3.5 plus, but qwen-plus is the known name for Alibaba DashScope
    model = "qwen3.5-plus"
    
    prompt = (
        "คุณเป็นนักเขียนนิยายระดับแนวหน้า จงเขียนนิยายไซไฟ (Sci-Fi) ธีมอวกาศและทะเลดวงดาว เรื่อง 'แมวทะเล กับ ผู้เฒ่าเต่า' "
        "ความยาวตลอดเรื่องสเกลใหญ่ประมาณ 10000 คำ ถ่ายทอดมิตรภาพต่างเผ่าพันธุ์ และการผจญภัยสุดล้ำ "
        "โปรดเน้นคำสำคัญด้วย **ตัวหนา** หรือ *ตัวเอียง* อย่างเหมาะสม เพื่อให้จัดรูปแบบสวยงาม "
        "ในข้อความแรกนี้ ให้เริ่มบรรยายบทนำและเนื้อเรื่องส่วนที่ 1 ก่อน แล้วเดี๋ยวจะให้เขียนเรื่องส่วนถัดไปเรื่อยๆ จนจบ"
    )
    
    messages = [{"role": "user", "content": prompt}]
    max_loops = 10
    full_text = ""
    total_tokens = 0
    total_time = 0
    
    console.print(f"[bold cyan]🚀 เริ่มมหกรรมเขียนนิยายต่อเนื่องด้วย Alibaba Cloud ({model})...[/bold cyan]")
    
    for loop in range(max_loops):
        console.print(f"\n[bold yellow]📖 เริ่มดึงเนื้อหาส่วนที่ {loop+1}/{max_loops}...[/bold yellow]\n")
        try:
            start_time = time.time()
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=2000, # DashScope max single response is slightly different
                stream=True,
                temperature=0.7
            )
            
            chunk_text = ""
            first_token_time = None
            
            for chunk in response:
                if chunk.choices and chunk.choices[0].delta.content:
                    if first_token_time is None:
                        first_token_time = time.time()
                    
                    text = chunk.choices[0].delta.content
                    chunk_text += text
                    full_text += text
                    total_tokens += 1
                    print(text, end="", flush=True)
                    
            print("\n")
            loop_gen_time = time.time() - start_time
            total_time += loop_gen_time
            
            messages.append({"role": "assistant", "content": chunk_text})
            
            if len(full_text) > 40000:
                break
                
            if loop == max_loops - 2:
                messages.append({"role": "user", "content": "จงเขียนสรุปบทจบของนิยายเรื่องนี้อย่างงดงามบริบูรณ์ในย่อหน้าถัดไป"})
            else:
                messages.append({"role": "user", "content": "เขียนเนื้อเรื่องต่อจากประโยคสุดท้ายแบบติดกันเลยให้ลื่นไหล"})
            
            time.sleep(1)
            
        except Exception as e:
            console.print(f"\n[bold red]❌ เกิดข้อผิดพลาด: {e}[/bold red]")
            break
            
    create_docx(full_text)
    
    console.print("="*60)
    console.print(f"[bold magenta]⏱️ Total Gen Time: {total_time:.2f} s[/bold magenta]")
    console.print(f"[bold magenta]📚 Total Tokens (Approx Chunks): {total_tokens}[/bold magenta]")
    console.print("="*60)

if __name__ == "__main__":
    run_novel_generation()
